#!/usr/bin/env python3
"""
Test script to create sample Hindi and Kannada documents and test end-to-end processing
"""
import sys
import asyncio
from pathlib import Path
from docx import Document

# Add src directory to Python path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from src.core.document_processor import DocumentProcessor
from src.models.ai_models import AIModelManager, DocumentAnalyzer

def create_test_documents():
    """Create test DOCX files with Hindi and Kannada content"""
    
    # Create Hindi test document
    hindi_doc = Document()
    hindi_doc.add_heading('हिंदी परीक्षण दस्तावेज़', 0)
    
    hindi_content = [
        'यह एक हिंदी भाषा का परीक्षण दस्तावेज़ है। इसमें भारत की राजधानी दिल्ली के बारे में जानकारी है।',
        'भारत एक विविधताओं से भरा देश है। यहाँ अनेक भाषाएँ बोली जाती हैं।',
        'हिंदी भारत की राजभाषा है और यह देवनागरी लिपि में लिखी जाती है।',
        'शिक्षा का महत्व: शिक्षा व्यक्तित्व विकास में महत्वपूर्ण भूमिका निभाती है।',
        'प्रौद्योगिकी के क्षेत्र में भारत ने अभूतपूर्व प्रगति की है।'
    ]
    
    for content in hindi_content:
        hindi_doc.add_paragraph(content)
    
    hindi_doc.save('test_hindi.docx')
    print("✅ Created test_hindi.docx")
    
    # Create Kannada test document
    kannada_doc = Document()
    kannada_doc.add_heading('ಕನ್ನಡ ಪರೀಕ್ಷಾ ದಾಖಲೆ', 0)
    
    kannada_content = [
        'ಇದು ಕನ್ನಡ ಭಾಷೆಯ ಪರೀಕ್ಷಾ ದಾಖಲೆಯಾಗಿದೆ। ಇಲ್ಲಿ ಕರ್ನಾಟಕ ರಾಜ್ಯದ ಬಗ್ಗೆ ಮಾಹಿತಿ ಇದೆ।',
        'ಕರ್ನಾಟಕ ದಕ್ಷಿಣ ಭಾರತದ ಪ್ರಮುಖ ರಾಜ್ಯವಾಗಿದೆ। ಬೆಂಗಳೂರು ಇದರ ರಾಜಧಾನಿಯಾಗಿದೆ।',
        'ಕನ್ನಡ ಒಂದು ದ್ರಾವಿಡ ಭಾಷೆಯಾಗಿದೆ ಮತ್ತು ಇದು ಕನ್ನಡ ಲಿಪಿಯಲ್ಲಿ ಬರೆಯಲಾಗುತ್ತದೆ।',
        'ಶಿಕ್ಷಣದ ಮಹತ್ವ: ಶಿಕ್ಷಣವು ವ್ಯಕ್ತಿತ್ವ ಅಭಿವೃದ್ಧಿಯಲ್ಲಿ ಪ್ರಮುಖ ಪಾತ್ರ ವಹಿಸುತ್ತದೆ।',
        'ತಂತ್ರಜ್ಞಾನ ಕ್ಷೇತ್ರದಲ್ಲಿ ಭಾರತವು ಅಪಾರ ಪ್ರಗತಿ ಸಾಧಿಸಿದೆ।'
    ]
    
    for content in kannada_content:
        kannada_doc.add_paragraph(content)
    
    kannada_doc.save('test_kannada.docx')
    print("✅ Created test_kannada.docx")
    
    # Create mixed language document
    mixed_doc = Document()
    mixed_doc.add_heading('Multi-Language Test Document / बहुभाषी परीक्षण', 0)
    
    mixed_content = [
        'This document contains multiple Indian languages.',
        'हिंदी: यह हिंदी भाषा का वाक्य है।',
        'ಕನ್ನಡ: ಇದು ಕನ್ನಡ ಭಾಷೆಯ ವಾಕ್ಯವಾಗಿದೆ।',
        'English: This demonstrates multilingual document processing.',
        'संस्कृत श्लोक: सर्वे भवन्तु सुखिनः सर्वे सन्तु निरामयाः।'
    ]
    
    for content in mixed_content:
        mixed_doc.add_paragraph(content)
    
    mixed_doc.save('test_mixed.docx')
    print("✅ Created test_mixed.docx")

async def test_document_processing():
    """Test the complete document processing pipeline"""
    print("\n" + "=" * 60)
    print("TESTING DOCUMENT PROCESSING PIPELINE")
    print("=" * 60)
    
    try:
        # Initialize document processor
        processor = DocumentProcessor()
        print("✅ DocumentProcessor initialized")
        
        # Test files
        test_files = [
            'test_hindi.docx',
            'test_kannada.docx', 
            'test_mixed.docx'
        ]
        
        results = {}
        
        for filename in test_files:
            if Path(filename).exists():
                print(f"\n📄 Processing {filename}...")
                
                try:
                    # Process document
                    processed_doc = await processor.process_document(filename)
                    
                    # Get statistics
                    stats = processor.get_document_stats(processed_doc)
                    
                    print(f"   ✅ Success: {len(processed_doc.elements)} elements extracted")
                    print(f"   📊 Languages detected: {stats['languages']}")
                    print(f"   📈 Element types: {stats['element_types']}")
                    print(f"   💯 Average confidence: {stats['avg_confidence']:.2f}")
                    
                    # Store results for summary testing
                    results[filename] = processed_doc
                    
                    # Print first few elements for verification
                    print(f"   📝 Sample elements:")
                    for i, elem in enumerate(processed_doc.elements[:3]):
                        print(f"      {i+1}. [{elem.language}] {elem.text[:100]}...")
                    
                except Exception as e:
                    print(f"   ❌ Error: {e}")
                    
            else:
                print(f"❌ File not found: {filename}")
        
        return results
        
    except Exception as e:
        print(f"❌ Failed to initialize processor: {e}")
        return {}

async def test_summarization(processed_docs):
    """Test AI summarization for Indian language documents"""
    print("\n" + "=" * 60)
    print("TESTING AI SUMMARIZATION")
    print("=" * 60)
    
    try:
        # Initialize AI models
        print("🧠 Initializing AI models...")
        ai_models = AIModelManager()
        analyzer = DocumentAnalyzer(ai_models)
        print("✅ AI models initialized")
        
        for filename, processed_doc in processed_docs.items():
            print(f"\n📝 Generating summary for {filename}...")
            
            try:
                # Generate summary
                summary_result = await analyzer.generate_document_summary(
                    processed_doc.elements,
                    summary_length='medium',
                    dual_language=True
                )
                
                print(f"   ✅ Summary generated successfully!")
                
                if isinstance(summary_result, dict):
                    if 'summary' in summary_result:
                        print(f"   📋 Summary: {summary_result['summary'][:200]}...")
                    if 'english_summary' in summary_result:
                        print(f"   🇬🇧 English: {summary_result['english_summary'][:200]}...")
                    if 'original_language' in summary_result:
                        print(f"   🌐 Language: {summary_result['original_language']}")
                else:
                    print(f"   📋 Summary: {summary_result[:200]}...")
                    
            except Exception as e:
                print(f"   ❌ Summarization failed: {e}")
    
    except Exception as e:
        print(f"❌ Failed to initialize AI models: {e}")

async def main():
    """Main test function"""
    print("🧪 POLYDOC END-TO-END TESTING")
    print("Testing Hindi and Kannada document processing and summarization")
    
    # Create test documents
    print("\n📁 Creating test documents...")
    create_test_documents()
    
    # Test document processing
    processed_docs = await test_document_processing()
    
    # Test summarization if we have processed documents
    if processed_docs:
        await test_summarization(processed_docs)
    
    print("\n" + "=" * 60)
    print("✅ END-TO-END TESTING COMPLETED")
    print("=" * 60)

if __name__ == "__main__":
    asyncio.run(main())
